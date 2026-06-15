"""
Script tạo file DOCX changelog cho Quiz Processor v1.5.0
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- HEADER BANNER ---
header_para = doc.add_paragraph()
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header_para.add_run('QUIZ PROCESSOR')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x5C, 0x35, 0xCC)  # purple

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run('Release Notes — Phiên bản 1.5.0')
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_para.add_run('Ngày phát hành: Tháng 6 / 2026')
run3.font.size = Pt(10)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# --- SECTION 1: Tổng quan ---
def add_section_title(doc, text, color_hex=(0x5C, 0x35, 0xCC)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color_hex)
    return p

def add_feature_block(doc, icon, title, items, bg='F5F3FF', border_color='7C3AED'):
    """Add a bordered feature block"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    
    # Title row
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run(f'{icon}  {title}')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x5C, 0x35, 0xCC)
    
    for item in items:
        p = cell.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if isinstance(item, tuple):
            # (text, bold_prefix)
            bold_text, rest_text = item
            r1 = p.add_run(bold_text)
            r1.font.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(rest_text)
            r2.font.size = Pt(10)
        else:
            r = p.add_run(item)
            r.font.size = Pt(10)
    
    last_p = cell.add_paragraph()
    last_p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

add_section_title(doc, '1.  Tính năng mới (New Features)')

add_feature_block(doc, '🔍', 'Luồng Kiểm định Cấu trúc Câu hỏi Độc lập', [
    ('Lệnh CLI mới --action doublecheck:', ' Chạy kiểm định cấu trúc câu hỏi hoàn toàn tách biệt với bước số hóa.'),
    ('Phát hiện 3 loại lỗi:', ''),
    '    •  option_count — Câu hỏi có ít hơn 3 hoặc nhiều hơn 5 đáp án.',
    '    •  stuck_options — Nhãn A./B./C. bị lẫn trong nội dung câu hỏi.',
    '    •  empty_option — Nhãn đáp án tồn tại nhưng không có nội dung text theo sau.',
    'Log chi tiết đầy đủ: số câu, nội dung câu hỏi, các đáp án nhận diện được.',
    'Nút "Kiểm định cấu trúc câu" xuất hiện cạnh nút "Bắt đầu Tạo đề" trong màn hình Tạo đề.',
], bg='F0FDF4', border_color='16A34A')

add_feature_block(doc, '📋', 'Feedback Loop Registry (feedback_loop.json)', [
    ('File registry JSON mới:', ' tự động ghi nhận lỗi cấu trúc câu sau mỗi lần kiểm định.'),
    ('Hai nguồn ghi nhận lỗi:', ''),
    '    •  auto — Ghi tự động khi chạy --action doublecheck.',
    '    •  manual — Ghi thủ công qua UI (card "Báo lỗi câu hỏi thủ công") hoặc --action add_feedback.',
    ('Cơ chế deduplicate:', ' Kiểm tra sự tồn tại trước khi ghi, tránh trùng lặp cùng một lỗi.'),
    ('Tự động làm sạch:', ' Khi chạy lại double-check, lỗi auto cũ được thay thế bởi kết quả mới; lỗi manual được giữ nguyên.'),
], bg='FFFBEB', border_color='D97706')

add_feature_block(doc, '🛠️', 'Sửa đổi File DOCX Nguồn từ Feedback', [
    ('Script scratch/fix_docx.py:', ' Vá lỗi cấu trúc file DOCX nguồn bằng cách đối chiếu với PDF gốc.'),
    'Khôi phục đáp án trống (A./B./C. không có text) từ nội dung PDF thực tế.',
    'Chèn thêm đáp án còn thiếu (D., E.) vào đúng vị trí bằng cách clone paragraph.',
    'Xóa các đoạn watermark nhiễu (ví dụ: "messages.downloaded_by") khỏi file DOCX.',
    ('Tìm kiếm theo nội dung:', ' Dùng text anchoring thay vì index cứng để tránh lệch khi có chèn/xóa đoạn.'),
    'Tự kiểm tra lại sau khi sửa và in log trạng thái chi tiết cho từng câu.',
], bg='EFF6FF', border_color='3B82F6')

add_feature_block(doc, '⌨️', 'Cập nhật Phím tắt Tùy biến', [
    ('Default key mapping mới:', ' A=1, B=2, C=3, D=5, E=8, Flag=6 (phản ánh đúng layout nút vật lý).'),
    ('Phím reserved:', ' Phím 4 = Gạch bỏ đáp án (Elimination), Phím 8 = Chuyển chế độ xem cuộn.'),
    ('Nút "Đặt lại mặc định":', ' Xuất hiện trong Settings > Phím tắt bàn phím, reset về bộ phím chuẩn.'),
    'Cảnh báo rõ trong UI: "Phím 4 và 8 đã được dùng cho chức năng khác, tránh dùng."',
], bg='FDF4FF', border_color='A855F7')

# --- SECTION 2: Cải tiến ---
add_section_title(doc, '2.  Cải tiến (Improvements)')

improvements = [
    ('Ưu tiên 2 mới trong heuristic phát hiện đáp án:', ' Highlight (bôi vàng) được đưa lên Ưu tiên 2, vượt qua Bold và Color anomaly. Chính xác nhất với các file DOCX có đánh dấu màu.'),
    ('Mô hình dữ liệu LineData mở rộng:', ' Thêm field is_highlighted để lưu trữ trạng thái highlight từ bước parse.'),
    ('Lọc watermark tăng cường:', ' patterns.py nhận diện và loại bỏ thêm các dạng watermark mới (lOMoARcPSD|..., messages.downloaded_by).'),
    ('Validation engine phân tách luồng:', ' Luồng kiểm định cấu trúc (double-check) độc lập với luồng số hóa, cho phép kiểm tra lại bất kỳ lúc nào.'),
    ('Log console trong Generate Screen:', ' Bố cục 2 cột — cột cấu hình (trái) và Log Console thời gian thực (phải, SelectionArea).'),
]

for title, rest in improvements:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(rest)
    r2.font.size = Pt(10)

doc.add_paragraph()

# --- SECTION 3: Bảng lệnh CLI mới ---
add_section_title(doc, '3.  Lệnh CLI mới trong v1.5.0')

tbl = doc.add_table(rows=3, cols=3)
tbl.style = 'Table Grid'

headers = ['Lệnh (--action)', 'Mô tả', 'Tham số chính']
header_row = tbl.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    set_cell_bg(cell, '5C35CC')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)

data = [
    ('doublecheck', 'Kiểm định cấu trúc câu hỏi độc lập, phát hiện 3 loại lỗi, ghi vào feedback_loop.json', '--answer-file  --workspace'),
    ('add_feedback', 'Thêm báo lỗi thủ công vào feedback registry', '--answer-file  --workspace  --question-index  --message'),
]

for row_idx, (action, desc, params) in enumerate(data, start=1):
    row = tbl.rows[row_idx]
    set_cell_bg(row.cells[0], 'EDE9FE')
    r0 = row.cells[0].paragraphs[0].add_run(action)
    r0.font.bold = True
    r0.font.name = 'Consolas'
    r0.font.size = Pt(9)

    r1 = row.cells[1].paragraphs[0].add_run(desc)
    r1.font.size = Pt(9)

    r2 = row.cells[2].paragraphs[0].add_run(params)
    r2.font.name = 'Consolas'
    r2.font.size = Pt(8)

doc.add_paragraph()

# --- SECTION 4: Thông tin phiên bản ---
add_section_title(doc, '4.  Thông tin Phiên bản')

version_info = [
    ('Phiên bản:', ' 1.5.0 (Build +5)'),
    ('pubspec.yaml:', ' version: 1.5.0+5'),
    ('Ngày phát hành:', ' Tháng 6 / 2026'),
    ('Tương thích:', ' Windows 10/11 (x64). Flutter 3.x stable. Python 3.10+.'),
    ('Phụ thuộc mới:', ' Không có thư viện mới — chỉ cải tiến logic nội bộ.'),
]

for title, rest in version_info:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(rest)
    r2.font.size = Pt(10)

doc.add_paragraph()

# --- FOOTER NOTE ---
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Quiz Processor — Tài liệu nội bộ — v1.5.0')
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# --- SAVE ---
out_path = r'd:\My_projects\Random_Essential\Quiz_Processor\docs\CHANGELOG_v1.5.0.docx'
doc.save(out_path)
print(f'[OK] Saved: {out_path}')
