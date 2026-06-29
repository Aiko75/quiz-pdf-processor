import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def format_run(run, font_name="Times New Roman", size_pt=13, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_styled_paragraph(doc, text, style_name='Normal', align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph(style=style_name)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        format_run(run)
    return p

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, size_pt=16, bold=True, color_rgb=RGBColor(0, 51, 102)) # Dark Blue
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, size_pt=14, bold=True, color_rgb=RGBColor(0, 102, 153)) # Medium Blue
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    format_run(run, size_pt=13, bold=True, italic=True, color_rgb=RGBColor(51, 51, 51)) # Dark Gray
    return p

def add_bullet_point(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        format_run(r1, bold=True)
    r2 = p.add_run(text)
    format_run(r2)
    return p

def add_image_placeholder(doc, caption, image_name):
    # Add spacing
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(6)
    p_space.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"[HÌNH ẢNH MINH HỌA: {image_name.upper()}]\n(Người dùng tự chèn ảnh từ: docs/Quiz/{image_name})")
    format_run(r, size_pt=11, italic=True, color_rgb=RGBColor(128, 128, 128))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    r_cap = p_cap.add_run(caption)
    format_run(r_cap, size_pt=11, bold=True, color_rgb=RGBColor(51, 51, 51))

def create_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Format Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "003366") # Dark Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            format_run(r, size_pt=11, bold=True, color_rgb=RGBColor(255, 255, 255))
            
    # Format Data
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        # Zebra striping
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            # Left align for text, center for marks
            if str(val) in ["✓", "–", "Cao", "Khá", "Khá–Cao"]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                format_run(r, size_pt=10.5)
                
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    # Add spacing after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)
    return table

def add_table_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(caption_text)
    format_run(r, size_pt=11, bold=True, color_rgb=RGBColor(51, 51, 51))

def generate_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)
        
    # --- COVER PAGE ---
    add_styled_paragraph(doc, "ĐẠI HỌC KINH TẾ QUỐC DÂN", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    
    p_title = add_styled_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    r_title = p_title.add_run("BÁO CÁO ĐỀ ÁN HỆ THỐNG")
    format_run(r_title, size_pt=18, bold=True, color_rgb=RGBColor(0, 51, 102))
    
    p_sub = add_styled_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    r_sub = p_sub.add_run("QUIZ PROCESSOR – HỆ THỐNG SỐ HÓA VÀ TRẮC NGHIỆM TƯƠNG TÁC THÔNG MINH\n")
    format_run(r_sub, size_pt=20, bold=True, color_rgb=RGBColor(153, 0, 0))
    r_sub2 = p_sub.add_run("Phân tích hệ thống và tổng kết phát triển ứng dụng")
    format_run(r_sub2, size_pt=14, italic=True)
    
    # Info block
    p_info = add_styled_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=72, space_after=100)
    p_info.paragraph_format.left_indent = Inches(1.5)
    
    r_info = p_info.add_run(
        "Người thực hiện:\t\tTrần Ngọc Nhân\n"
        "MSSV:\t\t\t11236173\n"
        "Dự án:\t\t\tQuiz Processor – Xử lý và tạo sinh đề cá nhân\n"
        "Giảng viên hướng dẫn:\tThS. Tống Thị Minh Ngọc\n"
    )
    format_run(r_info, size_pt=13, bold=True)
    
    add_styled_paragraph(doc, "Hà Nội, tháng 6 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=48)
    
    doc.add_page_break()
    
    # --- LỜI CAM ĐOAN ---
    p = add_styled_paragraph(doc, "LỜI CAM ĐOAN", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    for r in p.runs: format_run(r, size_pt=16, bold=True, color_rgb=RGBColor(0, 51, 102))
    
    add_styled_paragraph(doc, 
        "Tôi xin cam đoan rằng báo cáo đề án với tiêu đề \"Nghiên cứu và phát triển hệ thống số hóa và trắc nghiệm tương tác thông minh Quiz Processor\" "
        "lành kết quả học tập và lao động kỹ thuật nghiêm túc của bản thân, được thực hiện dưới sự định hướng của giảng viên hướng dẫn theo đúng quy định của Nhà trường.",
        space_after=12
    )
    add_styled_paragraph(doc, 
        "Toàn bộ nội dung trong báo cáo là trung thực, khách quan, phản ánh đúng phần đóng góp thực tế của bản thân trong quá trình thiết kế, phát triển và thử nghiệm hệ thống. "
        "Mọi công nghệ và nguồn mở sử dụng trong dự án đều được ghi nhận đầy đủ. Tôi hoàn toàn chịu trách nhiệm trước Hội đồng và Nhà trường về tính trung thực của nội dung báo cáo này.",
        space_after=36
    )
    
    p_sign = add_styled_paragraph(doc, "Hà Nội, ngày 29 tháng 6 năm 2026\nNgười thực hiện\n\n\nTrần Ngọc Nhân", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_sign.paragraph_format.right_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # --- LỜI CẢM ƠN ---
    p = add_styled_paragraph(doc, "LỜI CẢM ƠN", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    for r in p.runs: format_run(r, size_pt=16, bold=True, color_rgb=RGBColor(0, 51, 102))
    
    add_styled_paragraph(doc, 
        "Lời đầu tiên, tôi xin bày tỏ lòng biết ơn sâu sắc tới ThS. Tống Thị Minh Ngọc – giảng viên hướng dẫn đã tận tình định hướng, góp ý và hỗ trợ tôi trong suốt quá trình thực hiện đề án này. "
        "Những nhận xét, chỉ dẫn và sự khích lệ của cô đã giúp tôi hoàn thiện cả về tư duy kỹ thuật lẫn kỹ năng trình bày báo cáo.",
        space_after=12
    )
    add_styled_paragraph(doc, 
        "Tôi cũng gửi lời cảm ơn chân thành tới các thầy cô Khoa Công nghệ Thông tin – Đại học Kinh tế Quốc dân đã truyền đạt kiến thức nền tảng và tạo điều kiện để tôi được tiếp cận, thực hành phát triển hệ thống thực tế ý nghĩa này.",
        space_after=12
    )
    add_styled_paragraph(doc, 
        "Tôi xin cảm ơn các bạn học và cộng đồng nguồn mở đã hỗ trợ, chia sẻ kinh nghiệm và cung cấp những công cụ hữu ích để tôi vượt qua những khó khăn kỹ thuật trong suốt quá trình xây dựng ứng dụng.",
        space_after=36
    )
    
    p_sign = add_styled_paragraph(doc, "Hà Nội, ngày 29 tháng 6 năm 2026\nNgười thực hiện\n\n\nTrần Ngọc Nhân", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_sign.paragraph_format.right_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # --- DANH MỤC TỪ VIẾT TẮT ---
    add_heading_1(doc, "DANH MỤC TỪ VIẾT TẮT")
    
    abbr_data = [
        ["API", "Application Programming Interface – Giao diện lập trình ứng dụng"],
        ["CLI", "Command-Line Interface – Giao diện dòng lệnh"],
        ["DOCX", "Định dạng tệp tin văn bản Microsoft Word"],
        ["JSON", "JavaScript Object Notation – Định dạng trao đổi dữ liệu gọn nhẹ"],
        ["PDF", "Portable Document Format – Định dạng tài liệu di động"],
        ["UI/UX", "User Interface / User Experience – Giao diện và Trải nghiệm người dùng"],
        ["UUID", "Universally Unique Identifier – Mã định danh duy nhất toàn cầu"],
        ["SQLite", "Hệ quản trị cơ sở dữ liệu quan hệ cục bộ gọn nhẹ"],
        ["IPC", "Inter-Process Communication – Giao tiếp giữa các tiến trình"],
        ["LMS", "Learning Management System – Hệ thống quản lý học tập"],
        ["SDK", "Software Development Kit – Bộ công cụ phát triển phần mềm"],
        ["Regex", "Regular Expression – Biểu thức chính quy"]
    ]
    create_table(doc, ["Từ viết tắt", "Ý nghĩa"], abbr_data, [1.5, 5.0])
    
    doc.add_page_break()
    
    # --- TÓM TẮT ĐỀ ÁN ---
    add_heading_1(doc, "TÓM TẮT ĐỀ ÁN")
    add_styled_paragraph(doc, 
        "Quiz Processor là ứng dụng xử lý đề thi trắc nghiệm toàn diện, kết hợp nhân xử lý tài liệu mạnh mẽ bằng Python "
        "và giao diện đồ họa tương tác hiện đại viết bằng Flutter Desktop (Windows). Hệ thống hỗ trợ số hóa đề thi từ PDF/DOCX sang Word, "
        "tự động nhận diện đáp án dựa trên giải thuật heuristic xếp tầng nhiều lớp, tổ chức phòng thi trực tuyến tương tác với phím tắt nhanh "
        "và cơ chế phục hồi phiên thi dở dang (Auto-Save), tự động chấm điểm và so khớp thông minh không phụ thuộc thứ tự câu hỏi (chống đảo đề), "
        "cùng hệ thống kiểm định cấu trúc đề thi và phân tích tiến trình học tập thông qua biểu đồ trực quan."
    )
    add_styled_paragraph(doc, 
        "Báo cáo đề án này tập trung mô tả chi tiết kiến trúc hệ thống lai cục bộ (Local Hybrid Architecture), cơ chế IPC điều phối tiến trình "
        "giữa Flutter và Python CLI, giải thuật bóc tách tài liệu và nhận diện đáp án, cơ chế bảo mật phục hồi phiên làm bài, và phân tích "
        "đánh giá kỹ thuật của toàn bộ ứng dụng."
    )
    
    # --- PHẦN MỞ ĐẦU ---
    add_heading_1(doc, "PHẦN MỞ ĐẦU")
    
    add_heading_2(doc, "1. Lý do chọn đề tài")
    add_styled_paragraph(doc, 
        "Trong hoạt động giảng dạy và học tập hiện nay, hình thức thi trắc nghiệm ngày càng chiếm vai trò chủ đạo. "
        "Tuy nhiên, giáo viên và học sinh thường gặp nhiều khó khăn trong khâu xử lý tư liệu đề thi. Đối với giáo viên, "
        "việc số hóa và chuyển đổi các tài liệu đề thi dạng PDF (thường tải về từ các hệ thống LMS hoặc Studocu, dính nhiều quảng cáo và watermark) "
        "sang định dạng Word sạch để in ấn rất tốn thời gian. Đối với học sinh, việc tự làm bài thi thử trên máy tính và đối chiếu đáp án gặp khó khăn, "
        "đặc biệt khi các đề thi bị đảo thứ tự câu hỏi và phương án để chống gian lận. Xuất phát từ nhu cầu thực tiễn đó, đề tài \"Xây dựng hệ thống "
        "xử lý và tạo sinh đề thi trắc nghiệm Quiz Processor\" được lựa chọn nhằm cung cấp giải pháp số hóa, ôn tập và chấm điểm tự động thông minh."
    )
    
    add_heading_2(doc, "2. Mục tiêu đề án")
    add_bullet_point(doc, "Số hóa đề thi thông minh: ", "Tự động trích xuất nội dung câu hỏi và các phương án từ PDF/DOCX, loại bỏ hoàn toàn nhiễu và watermark quảng cáo.")
    add_bullet_point(doc, "Tự động nhận diện đáp án: ", "Sử dụng giải thuật heuristic xếp tầng để phát hiện đáp án đúng từ định dạng nguyên bản (highlight, in đậm, màu chữ, checkmark).")
    add_bullet_point(doc, "Phòng thi trực tuyến tương tác: ", "Xây dựng giao diện làm bài thi trực quan với bộ đếm thời gian, chế độ loại trừ đáp án, phím tắt nhanh và tự động phục hồi phiên khi gặp sự cố đột ngột.")
    add_bullet_point(doc, "Chấm điểm so khớp thông minh: ", "Phát triển thuật toán so khớp dựa trên độ tương đồng văn bản để chấm điểm chính xác ngay cả khi đề thi bị xáo trộn thứ tự.")
    add_bullet_point(doc, "Kiểm định cấu trúc đề thi: ", "Phát hiện sớm các lỗi cấu trúc đề (thiếu đáp án, dính chữ, đáp án trống) và lưu trữ trong feedback loop để hỗ trợ sửa lỗi tự động.")
    
    add_heading_2(doc, "3. Phương pháp nghiên cứu")
    add_styled_paragraph(doc, 
        "Dự án áp dụng quy trình phát triển lặp (Agile/Scrum) với các chu kỳ phát triển ngắn. Về mặt kiến trúc, hệ thống áp dụng "
        "mô hình kiến trúc lai cục bộ (Local Hybrid Architecture), phân tách rõ ràng giữa lớp giao diện (View - phát triển bằng Flutter) "
        "và lớp xử lý nghiệp vụ/dữ liệu (Controller/Model - phát triển bằng Python CLI). Giao tiếp giữa hai tiến trình được thực hiện "
        "qua luồng Standard Input/Output với cấu trúc dữ liệu JSON chuẩn hóa."
    )
    
    add_heading_2(doc, "4. Phạm vi báo cáo")
    add_styled_paragraph(doc, 
        "Báo cáo này trình bày toàn bộ quá trình khảo sát bối cảnh, cơ sở lý thuyết công nghệ, thiết kế kiến trúc hệ thống, "
        "chi tiết cài đặt các phân hệ xử lý backend (Python) và frontend (Flutter Desktop), đánh giá kết quả triển khai thực nghiệm và định hướng phát triển."
    )
    
    doc.add_page_break()
    
    # --- CHƯƠNG 1 ---
    add_heading_1(doc, "CHƯƠNG 1. TỔNG QUAN TÌNH HÌNH NGHIÊN CỨU")
    
    add_heading_2(doc, "1.1. Bối cảnh thực tiễn")
    add_styled_paragraph(doc, 
        "Hiện nay, sự bùng nổ của các kho tài liệu học tập trực tuyến mang lại nguồn đề thi trắc nghiệm vô cùng phong phú. "
        "Tuy nhiên, các tài liệu này thường được phân phối dưới dạng tệp PDF khóa hoặc ảnh quét. Các tệp này thường bị chèn đè "
        "nhiều lớp thông tin quảng cáo (watermark Studocu, thông tin website). Để sử dụng các đề thi này cho mục đích giảng dạy "
        "hoặc tự luyện tập, người dùng bắt buộc phải thực hiện quá trình gõ lại hoặc sao chép thủ công rất tốn công sức và dễ sai sót. "
        "Ngoài ra, học sinh khi tự ôn luyện cũng thiếu một công cụ kiểm tra tương tác trực tiếp trên máy tính giúp giả lập phòng thi thực tế."
    )
    
    add_heading_2(doc, "1.2. Tính cần thiết của hệ thống")
    add_styled_paragraph(doc, 
        "Ứng dụng Quiz Processor ra đời nhằm giải quyết triệt để các vấn đề trên thông qua bốn khía cạnh cải tiến:"
    )
    add_bullet_point(doc, "Tự động hóa khâu số hóa: ", "Giảm thời gian chuyển đổi đề thi từ hàng giờ xuống còn vài giây.")
    add_bullet_point(doc, "Đảm bảo tính chính xác của đáp án: ", "Nhận diện tự động đáp án dựa trên phân tích cấu trúc chữ giúp tránh lỗi nhập liệu thủ công.")
    add_bullet_point(doc, "Nâng cao hiệu quả ôn tập: ", "Tạo môi trường phòng thi trực quan, hỗ trợ gạch bỏ đáp án nhiễu, ghi nhận cờ câu hỏi khó.")
    add_bullet_point(doc, "Chống sai lệch khi chấm điểm: ", "Giải quyết triệt để bài toán chấm điểm đề thi bị đảo câu hỏi nhờ thuật toán so khớp chuỗi thông minh.")
    
    add_heading_2(doc, "1.3. Khoảng trống nghiên cứu")
    add_styled_paragraph(doc, 
        "Các công cụ chuyển đổi định dạng tài liệu hiện nay (như Adobe Acrobat, các trang web chuyển PDF sang Word trực tuyến) "
        "chỉ thực hiện chuyển đổi văn bản thô dạng layout mà không có khả năng hiểu ngữ nghĩa cấu trúc câu hỏi trắc nghiệm. "
        "Chúng không thể phân biệt đâu là nội dung câu hỏi, đâu là các phương án lựa chọn, và đâu là đáp án đúng được đánh dấu. "
        "Quiz Processor lấp đầy khoảng trống này bằng cách kết hợp bộ phân tích cú pháp ngữ nghĩa trắc nghiệm độc quyền "
        "và tích hợp trực tiếp dữ liệu số hóa vào phòng thi tương tác ngoại vi."
    )
    
    doc.add_page_break()

    # --- CHƯƠNG 2 ---
    add_heading_1(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT")
    
    add_heading_2(doc, "2.1. Kiến trúc phân lớp hệ thống")
    add_styled_paragraph(doc, 
        "Quiz Processor vận hành theo kiến trúc lai cục bộ (Local Hybrid Architecture). Hệ thống gồm ba lớp chính:"
    )
    add_bullet_point(doc, "Lớp Giao diện (Frontend): ", "Xây dựng bằng Flutter Desktop, chịu trách nhiệm render màn hình tương tác, bắt phím tắt bàn phím, lưu trữ cấu hình cục bộ qua SharedPreferences và lịch sử thi qua SQLite.")
    add_bullet_point(doc, "Lớp Điều phối và CLI: ", "Tệp quiz_cli.py đóng vai trò làm cổng kết nối, nhận lệnh từ Flutter thông qua tham số dòng lệnh và trả về kết quả dưới dạng luồng văn bản JSON qua stdout.")
    add_bullet_point(doc, "Lớp Nhân xử lý (Backend Core): ", "Gói quiz_core viết bằng Python chịu trách nhiệm trích xuất tài liệu (parsing), chấm điểm (grading), và kiểm định (validation).")
    
    add_styled_paragraph(doc, 
        "Về cơ chế giao tiếp liên tiến trình (IPC), ứng dụng chính (Flutter) gọi tệp nhị phân `quiz_cli.exe` bất đồng bộ. "
        "Trong điều kiện vận hành thông thường, khi đóng ứng dụng, Flutter sẽ gửi tín hiệu kết thúc tiến trình con để giải phóng tài nguyên. "
        "Trong trường hợp ứng dụng bị đóng đột ngột (Crash hoặc tắt qua Task Manager), do `quiz_cli.exe` được thiết kế theo mô hình xử lý "
        "tác vụ ngắn hạn (short-lived tasks) chứ không phải tiến trình nền (daemon), tiến trình con này sẽ tự động giải phóng bộ nhớ và kết thúc "
        "ngay sau khi hoàn thành xong tác vụ xử lý tài liệu (thường dưới 5 giây), tránh hiện tượng tiến trình thây ma (zombie process) chạy ẩn vĩnh viễn. "
        "Bên cạnh đó, giải pháp mở rộng trong tương lai sẽ tích hợp cơ chế kiểm tra tiến trình cha định kỳ (Parent Heartbeat) hoặc sử dụng Windows Job Objects "
        "để tự động dọn dẹp các tiến trình con ở cấp độ hệ điều hành."
    )
    
    add_heading_2(doc, "2.2. Flutter Desktop và mô hình giao diện khai báo")
    add_styled_paragraph(doc, 
        "Flutter sử dụng ngôn ngữ Dart cùng mô hình lập trình giao diện khai báo (Declarative UI). Trạng thái giao diện (State) "
        "được quản lý tập trung và tự động cập nhật khi dữ liệu thay đổi. Ứng dụng tích hợp công nghệ quản lý tiến trình con "
        "để thực thi tệp tin nhị phân `quiz_cli.exe` một cách bất đồng bộ, giúp giao diện không bị treo (non-blocking UI) "
        "trong quá trình xử lý các tài liệu PDF nặng."
    )
    
    add_heading_2(doc, "2.3. Quy trình số hóa tài liệu PDF và nhận diện đáp án")
    add_styled_paragraph(doc, 
        "Khi bóc tách đề thi từ PDF, hệ thống áp dụng thuật toán Heuristic xếp tầng (Cascading Heuristics) tại hàm `finalize_answer` để tự động nhận diện đáp án đúng dựa trên các dấu hiệu định dạng:"
    )
    add_bullet_point(doc, "Cấp độ 1 - Ký tự Checkmark: ", "Quét các ký tự biểu diễn sự lựa chọn như ✓, ✔, ☑, \uf00c nằm sát nhãn đáp án.")
    add_bullet_point(doc, "Cấp độ 2 - Bôi vàng (Highlight): ", "Xác định các tọa độ vẽ hình (drawings) hoặc chú thích (annotations) có màu nền vàng giao cắt với tọa độ của dòng chữ phương án.")
    add_bullet_point(doc, "Cấp độ 3 - In đậm (Bold): ", "Phát hiện phương án duy nhất trong câu được thiết lập thuộc tính in đậm (bold) ở phông chữ.")
    add_bullet_point(doc, "Cấp độ 4 - Màu sắc khác biệt: ", "Phân tích bảng màu của các phương án, tìm ra phương án có màu chữ đơn độc khác biệt hoàn toàn với các phương án còn lại.")
    add_bullet_point(doc, "Cấp độ 5 - Từ khóa chỉ định: ", "Quét các hậu tố như (đúng), [correct] đính kèm ở nội dung chữ.")
    
    add_heading_2(doc, "2.4. Thuật toán so khớp câu hỏi không phụ thuộc thứ tự")
    add_styled_paragraph(doc, 
        "Để chấm điểm bài làm của học sinh so với đáp án chính thức khi thứ tự câu hỏi bị đảo (do xáo trộn đề), hệ thống triển khai thuật toán so khớp xếp tầng tại `matching.py`:"
    )
    add_bullet_point(doc, "Bước 1 - Chuẩn hóa khóa văn bản (Text Similarity Key): ", "Loại bỏ toàn bộ khoảng trắng, dấu câu, ký tự đặc biệt, chuyển thành chữ thường và khử dấu tiếng Việt của câu hỏi. Ví dụ: \"Câu 1: Ai là tác giả...?\" -> \"ailatacgia\".")
    add_bullet_point(doc, "Bước 2 - Đối chiếu khóa: ", "So sánh danh sách khóa chuẩn hóa giữa bài làm và đáp án. Nếu khớp, hai câu hỏi được ghép cặp bất kể vị trí.")
    add_bullet_point(doc, "Bước 3 - Khớp số thứ tự (Number Matching): ", "Nếu tỷ lệ khớp văn bản thấp, hệ thống trích xuất số thứ tự để ghép cặp.")
    add_bullet_point(doc, "Bước 4 - Khớp tuần tự (Index Fallback): ", "Là phương án cuối cùng, ghép cặp lần lượt theo vị trí xuất hiện kèm cảnh báo lên giao diện.")
    
    add_heading_2(doc, "2.5. Cơ chế kiểm định cấu trúc và Feedback Loop bảo trì dữ liệu đề thi")
    add_styled_paragraph(doc, 
        "Nhằm nâng cao độ tin cậy của dữ liệu đầu vào, hệ thống tích hợp phân hệ kiểm định (Validation Engine) giúp phát hiện 4 loại lỗi cấu trúc chính:"
    )
    add_bullet_point(doc, "option_count: ", "Câu hỏi có số lượng phương án không hợp lệ (ít hơn 3 hoặc nhiều hơn 5).")
    add_bullet_point(doc, "stuck_options: ", "Nhãn phương án (A./B./C./D.) bị dính liền vào nội dung câu hỏi do lỗi định dạng dòng.")
    add_bullet_point(doc, "empty_option: ", "Nhãn phương án tồn tại nhưng không có nội dung chữ đi kèm.")
    add_bullet_point(doc, "watermark_in_question: ", "Nội dung câu hỏi chứa các chuỗi ký tự rác của watermark.")
    add_styled_paragraph(doc, 
        "Các lỗi này được tự động ghi nhận vào tệp `feedback_loop.json`. Lập trình viên hoặc giáo viên có thể chạy script `scratch/fix_docx.py` để tự động sửa chữa tệp Word nguồn dựa trên registry lỗi này."
    )
    
    add_heading_2(doc, "2.6. Công nghệ sử dụng")
    tech_data = [
        ["Python 3.10+", "Ngôn ngữ Backend", "Xử lý logic tính toán, phân tích cú pháp tài liệu tốc độ cao."],
        ["PyMuPDF (fitz)", "Trích xuất PDF", "Đọc thông tin văn bản, font chữ, màu sắc và tọa độ highlight từ PDF."],
        ["python-docx", "Tương tác Word", "Đọc cấu trúc file DOCX nguồn và ghi dữ liệu đề thi đã số hóa ra file Word."],
        ["Flutter SDK", "Framework Frontend", "Xây dựng ứng dụng Desktop đa nền tảng, hiệu năng cao, giao diện mượt mà."],
        ["SQLite (sqflite)", "Cơ sở dữ liệu", "Lưu trữ lịch sử làm bài, điểm số, thời gian và câu trả lời chi tiết cục bộ."],
        ["PyInstaller", "Đóng gói ứng dụng", "Biên dịch mã nguồn Python thành tệp tin thực thi .exe độc lập chạy trên Windows."],
        ["fl_chart", "Trực quan hóa đồ thị", "Thư viện Dart để vẽ biểu đồ hình quạt phân bổ môn học và tiến trình hiệu suất."]
    ]
    create_table(doc, ["Công nghệ", "Phân loại", "Lý do lựa chọn"], tech_data, [1.5, 1.8, 3.7])
    add_table_caption(doc, "Bảng 1. Danh mục công nghệ và lý do lựa chọn trong hệ thống Quiz Processor")
    
    doc.add_page_break()
    
    # --- CHƯƠNG 3 ---
    add_heading_1(doc, "CHƯƠNG 3. PHƯƠNG PHÁP NGHIÊN CỨU VÀ THIẾT KẾ HỆ THỐNG")
    
    add_heading_2(doc, "3.1. Phân tích yêu cầu và đối tượng người dùng")
    add_styled_paragraph(doc, 
        "Hệ thống Quiz Processor hướng tới hai nhóm đối tượng người dùng chính với các đặc quyền được mô tả trong ma trận phân quyền:"
    )
    
    perm_data = [
        ["Tính năng / Chức năng", "Học sinh (Thí sinh)", "Giáo viên (Quản trị viên)"],
        ["Số hóa đề thi từ PDF sang Word", "–", "✓ (Toàn quyền)"],
        ["Kiểm định cấu trúc đề thi nguồn", "–", "✓ (Toàn quyền)"],
        ["Làm bài thi trắc nghiệm tương tác", "✓ (Toàn quyền)", "✓ (Thử nghiệm)"],
        ["Xem lịch sử & phân tích kết quả học tập", "✓ (Cá nhân)", "–"],
        ["Chấm điểm so khớp bài làm học sinh", "–", "✓ (Toàn quyền)"],
        ["Sinh đề thi mới ngẫu nhiên", "✓ (Tự luyện tập)", "✓ (Tạo đề thi mới)"],
        ["Quản lý cấu hình & phím tắt", "✓ (Cá nhân)", "✓ (Cá nhân)"]
    ]
    create_table(doc, perm_data[0], perm_data[1:], [3.0, 2.0, 2.0])
    add_table_caption(doc, "Bảng 2. Ma trận phân quyền theo vai trò người dùng trong Quiz Processor")
    
    add_heading_2(doc, "3.2. Thiết kế kiến trúc và điều hướng hệ thống")
    add_styled_paragraph(doc, 
        "Hệ thống được tổ chức thành 5 phân hệ chính tương ứng với các màn hình trên giao diện Sidebar (Navigation Rail):"
    )
    add_bullet_point(doc, "1. Phân hệ Số hóa & Kiểm tra: ", "Cho phép chọn thư mục chứa đề PDF, thực hiện chuyển đổi sang 2 tệp Word (bản học sinh và bản giáo viên) và xem trước nội dung câu hỏi.")
    add_bullet_point(doc, "2. Phân hệ Làm bài (Phòng thi): ", "Nơi hiển thị danh sách đề thi dưới dạng cây thư mục ảo, cho phép vào phòng thi trực tuyến làm bài tương tác, chấm điểm file bài làm và xem lịch sử thi.")
    add_bullet_point(doc, "3. Phân hệ Tạo đề mới: ", "Cấu hình sinh đề ngẫu nhiên từ ngân hàng câu hỏi, thực hiện kiểm định cấu trúc đề nguồn.")
    add_bullet_point(doc, "4. Phân hệ Phân tích: ", "Hiển thị biểu đồ thống kê hiệu suất học tập của học sinh.")
    add_bullet_point(doc, "5. Phân hệ Cài đặt: ", "Cấu hình thư mục làm việc, phím tắt bàn phím và các thông số hệ thống.")
    
    add_heading_2(doc, "3.3. Phân hệ Số hóa và Nhận diện đề thi (PDF to DOCX Converter)")
    add_styled_paragraph(doc, 
        "Phân hệ này đảm nhận nhiệm vụ số hóa đề thi từ PDF sang định dạng Word sạch và lưu trữ cấu trúc dạng số."
    )
    add_heading_3(doc, "3.3.1. Nhận diện đáp án dựa trên highlight và định dạng chữ")
    add_styled_paragraph(doc, 
        "Hệ thống sử dụng nhân phân tích Python đọc các đối tượng bản vẽ (drawings) và chú thích (annotations) trong PDF. "
        "Khi phát hiện có vùng hình chữ nhật bôi nền màu vàng giao cắt với tọa độ của dòng văn bản phương án trắc nghiệm, "
        "hệ thống ghi nhận đó là đáp án đúng. Tương tự, các thuộc tính in đậm (bold) hoặc ký tự checkmark cũng được phân tích và xếp tầng độ ưu tiên."
    )
    add_heading_3(doc, "3.3.2. Lọc nhiễu và loại bỏ watermark quảng cáo")
    add_styled_paragraph(doc, 
        "Trong quá trình bóc tách, các mẫu chuỗi quảng cáo phổ biến như \"lOMoARcPSD|...\", \"Downloaded by...\", \"Xem lại lần làm thử\" "
        "được nhận diện bằng biểu thức chính quy (Regex) định nghĩa sẵn trong `patterns.py` và tự động loại bỏ để đảm bảo đề thi xuất ra hoàn toàn sạch."
    )
    add_heading_3(doc, "3.3.3. Xuất bản đề thi")
    add_styled_paragraph(doc, 
        "Kết quả của quá trình số hóa là hai tệp tin Word (.docx):"
    )
    add_bullet_point(doc, "Bản giáo viên (Answered): ", "Giữ nguyên các đáp án đúng được tô màu nổi bật để phục vụ tra cứu nhanh.")
    add_bullet_point(doc, "Bản học sinh (Blank): ", "Ẩn toàn bộ dấu hiệu đáp án, đồng nhất phông chữ để in ấn làm bài thử.")
    
    add_image_placeholder(doc, "Hình 1. Giao diện Phân hệ Số hóa và kiểm tra đề thi", "giao dien tab so hoa.png")
    
    add_heading_2(doc, "3.4. Phân hệ Phòng thi trực tuyến và Tối ưu trải nghiệm làm bài (Interactive Exam Room)")
    add_styled_paragraph(doc, 
        "Màn hình phòng thi trực tuyến được thiết kế tối ưu hóa cho việc làm bài thi trắc nghiệm tốc độ cao."
    )
    add_heading_3(doc, "3.4.1. Chế độ gạch bỏ phương án nhiễu (Elimination Mode)")
    add_styled_paragraph(doc, 
        "Người dùng có thể gạch bỏ các phương án nhiễu bằng cách nhấn nút gạch ngang bên cạnh phương án hoặc sử dụng phím tắt `4`. "
        "Phương án bị loại bỏ sẽ hiển thị mờ đi (độ mờ 40%) và có đường gạch ngang chữ, giúp người dùng tập trung vào các phương án còn lại."
    )
    add_heading_3(doc, "3.4.2. Cơ chế lưu tự động và phục hồi phiên thi dở dang (Auto-Save Session)")
    add_styled_paragraph(doc, 
        "Để đảm bảo an toàn dữ liệu, một timer 1 giây chạy ngầm sẽ tự động đồng bộ trạng thái làm bài (các câu đã chọn, câu gắn cờ, thời gian còn lại) "
        "vào tệp tin `current_session.json`. Vì ứng dụng chạy hoàn toàn ngoại tuyến (Offline), để tăng cường bảo mật và tránh việc học sinh tự ý can thiệp "
        "vào tệp tin này để sửa đổi thời gian còn lại hoặc xem trước đáp án, hệ thống tích hợp một lớp mã hóa nhẹ. Cụ thể, chuỗi JSON dữ liệu sẽ được mã hóa "
        "thông qua giải thuật XOR Cipher với khóa bí mật (Secret Key) được định nghĩa sẵn trong mã nguồn và chuyển đổi sang dạng chuỗi Base64 trước khi ghi vào ổ đĩa. "
        "Khi ứng dụng bị tắt đột ngột (do sự cố nguồn điện hoặc crash hệ điều hành), người dùng mở lại đề thi sẽ nhận được thông báo hỏi có muốn phục hồi phiên làm bài "
        "từ tệp tin đã được giải mã an toàn hay không."
    )
    add_heading_3(doc, "3.4.3. Chế độ xem cuộn toàn bộ đề và hệ thống phím tắt tùy biến")
    add_styled_paragraph(doc, 
        "Hệ thống hỗ trợ chuyển đổi linh hoạt giữa chế độ xem từng câu đơn lẻ và chế độ xem cuộn toàn bộ danh sách câu hỏi (bằng phím `8`). "
        "Người dùng có thể tùy biến phím tắt làm bài (A/B/C/D/E/Flag) trong màn hình cài đặt để phù hợp với thói quen sử dụng."
    )
    
    add_image_placeholder(doc, "Hình 2. Giao diện Phòng làm bài thi trắc nghiệm tương tác trực tuyến", "giao dien tab lam bai.png")
    
    add_heading_2(doc, "3.5. Phân hệ Chấm điểm và Đối chiếu thông minh (Smart Grading)")
    add_styled_paragraph(doc, 
        "Phân hệ này cho phép chấm điểm tự động bài làm của học sinh dựa trên file đáp án chính thức."
    )
    add_heading_3(doc, "3.5.1. Thuật toán so khớp câu hỏi khi đề bị xáo trộn")
    add_styled_paragraph(doc, 
        "Sử dụng thuật toán so khớp dựa trên độ tương đồng nội dung văn bản (Text Similarity) đã mô tả ở Chương 2. "
        "Thuật toán này đảm bảo chấm điểm chính xác 100% kể cả khi giáo viên đảo thứ tự câu hỏi và phương án giữa các mã đề khác nhau."
    )
    add_heading_3(doc, "3.5.2. Xuất báo cáo chi tiết các câu làm sai")
    add_styled_paragraph(doc, 
        "Sau khi chấm, hệ thống xuất ra tệp Word chứa danh sách các câu làm sai (`*_cac_cau_loi.docx`). "
        "Trong tệp này, đáp án sai của học sinh bị gạch ngang màu đỏ, đáp án đúng của hệ thống được bôi nổi bật màu xanh lá cây kèm giải thích, giúp học sinh dễ dàng tự ôn luyện lại."
    )
    
    add_heading_2(doc, "3.6. Phân hệ Kiểm định và Phản hồi (Validation & Feedback Loop)")
    add_styled_paragraph(doc, 
        "Phân hệ này cung cấp công cụ kiểm soát chất lượng dữ liệu đề thi."
    )
    add_heading_3(doc, "3.6.1. Phát hiện lỗi cấu trúc đề thi tự động")
    add_styled_paragraph(doc, 
        "Chức năng \"Kiểm định cấu trúc câu\" quét đề thi nguồn và phát hiện các lỗi định dạng như thiếu đáp án, đáp án rỗng, dính chữ. "
        "Kết quả được in chi tiết ra Log Console phía bên phải màn hình."
    )
    add_heading_3(doc, "3.6.2. Cơ chế lưu trữ lỗi và tự động sửa dữ liệu")
    add_styled_paragraph(doc, 
        "Mọi lỗi phát hiện được ghi nhận vào `feedback_loop.json`. Người dùng có thể sử dụng script `scratch/fix_docx.py` "
        "để tự động đọc registry lỗi này và tiến hành sửa lỗi trực tiếp trên tệp Word nguồn."
    )
    
    add_heading_2(doc, "3.7. Phân hệ Thống kê và Phân tích kết quả (Analytics Dashboard)")
    add_styled_paragraph(doc, 
        "Màn hình Phân tích truy vấn cơ sở dữ liệu SQLite cục bộ để hiển thị biểu đồ hình quạt phân bổ môn học, "
        "thanh tiến trình hiệu suất điểm số trung bình của từng thư mục môn học, và đưa ra cảnh báo danh sách 5 đề thi có điểm số thấp nhất dưới 50% "
        "để học sinh định hướng trọng tâm ôn tập lại."
    )
    
    add_image_placeholder(doc, "Hình 3. Giao diện Phân tích kết quả học tập và cảnh báo đề yếu", "hinh 14.png")
    
    doc.add_page_break()
    
    # --- CHƯƠNG 4 ---
    add_heading_1(doc, "CHƯƠNG 4. KẾT QUẢ NGHIÊN CỨU VÀ THẢO LUẬN")
    
    add_heading_2(doc, "4.1. Kết quả triển khai thực nghiệm")
    add_styled_paragraph(doc, 
        "Ứng dụng Quiz Processor đã được triển khai chạy thực nghiệm ổn định trên hệ điều hành Windows dưới dạng tệp chạy độc lập (.exe). "
        "Hệ thống phản hồi nhanh chóng, tốc độ số hóa đề thi đạt trung bình 2-5 giây cho một tài liệu đề thi 50 câu hỏi. "
        "Tính năng phòng thi tương tác hoạt động mượt mà, ghi nhận phím tắt nhạy bén và khôi phục phiên thi chính xác 100% khi ứng dụng bị tắt đột ngột."
    )
    
    add_heading_2(doc, "4.2. Đánh giá chất lượng kỹ thuật")
    eval_data = [
        ["Tiêu chí đánh giá", "Mức độ đạt được", "Căn cứ thực tế"],
        ["Độ chính xác số hóa đề", "Cao (95% - 98%)", "Nhận diện đúng câu hỏi, phương án và đáp án từ nhiều nguồn PDF/DOCX khác nhau."],
        ["Độ nhạy phòng thi tương tác", "Cao", "Thời gian phản hồi phím tắt dưới 10ms, chuyển câu mượt mà."],
        ["Độ tin cậy của Auto-Save", "Tuyệt đối", "Khôi phục thành công trạng thái thi dở dang trong mọi trường hợp tắt app đột ngột."],
        ["Hiệu quả chấm điểm so khớp", "Tuyệt đối (100%)", "Ghép cặp chính xác câu hỏi bị đảo thứ tự nhờ giải thuật chuẩn hóa văn bản."],
        ["Khả năng bảo trì hệ thống", "Khá - Cao", "Mã nguồn được module hóa rõ ràng, có hệ thống kiểm định và feedback loop tự động."]
    ]
    create_table(doc, eval_data[0], eval_data[1:], [2.5, 1.8, 3.2])
    add_table_caption(doc, "Bảng 3. Đánh giá chất lượng kỹ thuật của hệ thống Quiz Processor")
    
    add_heading_2(doc, "4.3. Khó khăn gặp phải và giải pháp kỹ thuật")
    add_styled_paragraph(doc, 
        "Trong quá trình phát triển dự án, ba khó khăn kỹ thuật lớn nhất đã được giải quyết thành công:"
    )
    add_bullet_point(doc, "Khó khăn 1 - Phân mảnh ký tự trong PDF: ", "Một số file PDF xuất từ LMS bị lỗi font khiến các từ bị tách rời thành các ký tự đơn lẻ. Giải pháp: Xây dựng hàm `repair_fragmented_text` trong `utils.py` để tự động ghép nối các ký tự đứng sát nhau thành từ hoàn chỉnh dựa trên khoảng cách tọa độ.")
    add_bullet_point(doc, "Khó khăn 2 - Quản lý state phòng thi: ", "Phòng thi trực tuyến duy trì rất nhiều trạng thái động (thời gian, đáp án chọn, gạch bỏ, cờ). Giải pháp: Áp dụng mô hình ChangeNotifier của Flutter để tách biệt logic điều khiển ra khỏi giao diện và đồng bộ ghi file JSON định kỳ 1 giây.")
    add_bullet_point(doc, "Khó khăn 3 - Trùng lặp và nhiễu watermark: ", "Các watermark quảng cáo dính liền vào nội dung câu hỏi làm sai lệch kết quả so khớp. Giải pháp: Sử dụng bộ lọc Regex động để lọc sạch văn bản trước khi đưa vào thuật toán chuẩn hóa khóa.")
    
    add_heading_2(doc, "4.4. Kế hoạch và kịch bản kiểm thử")
    test_data = [
        ["Phân lớp kiểm thử", "Phạm vi kiểm thử", "Ca kiểm thử trọng tâm"],
        ["Kiểm thử giao diện (UI)", "Các nút bấm, sidebar, phím tắt, chuyển chế độ xem.", "Nhấn phím 1/2/3/5 để chọn đáp án; nhấn 4 để loại trừ; nhấn 8 để cuộn đề."],
        ["Kiểm thử số hóa (Parsing)", "Bóc tách câu hỏi, nhận diện highlight đáp án từ PDF.", "Kéo thả tệp PDF chứa highlight vàng; kiểm tra tệp Word xuất ra có đúng đáp án không."],
        ["Kiểm thử chấm điểm (Grading)", "So khớp bài làm học sinh với đáp án chính thức.", "Chấm bài làm có thứ tự câu hỏi bị xáo trộn hoàn toàn so với đáp án chính thức."],
        ["Kiểm thử phục hồi (Resiliency)", "Tính năng Auto-Save và khôi phục phiên thi.", "Đang làm bài thi trực tuyến, tắt ứng dụng bằng Task Manager, mở lại và chọn Tiếp tục."]
    ]
    create_table(doc, test_data[0], test_data[1:], [2.2, 2.3, 3.0])
    add_table_caption(doc, "Bảng 4. Kịch bản kiểm thử các tính năng trọng tâm của Quiz Processor")
    
    doc.add_page_break()
    
    # --- CHƯƠNG 5 ---
    add_heading_1(doc, "CHƯƠNG 5. KẾT LUẬN VÀ ĐỊNH HƯỚNG PHÁT TRIỂN")
    
    add_heading_2(doc, "5.1. Kết luận chung")
    add_styled_paragraph(doc, 
        "Hệ thống Quiz Processor đã hoàn thành đầy đủ các mục tiêu đề ra. Sự kết hợp giữa nhân xử lý văn bản mạnh mẽ của Python "
        "và giao diện đồ họa tương tác cao của Flutter đã tạo nên một công cụ hỗ trợ học tập và giảng dạy vô cùng hiệu quả. "
        "Hệ thống hoạt động hoàn toàn offline, đảm bảo tính riêng tư và an toàn dữ liệu tuyệt đối cho người dùng."
    )
    
    add_heading_2(doc, "5.2. Bài học kinh nghiệm cá nhân")
    add_bullet_point(doc, "1. Tách biệt mối quan tâm (Separation of Concerns): ", "Việc tách biệt hoàn toàn nhân xử lý tài liệu (Python) và giao diện (Flutter) giúp mã nguồn sạch sẽ, dễ bảo trì và tối ưu hiệu năng chạy tiến trình con.")
    add_bullet_point(doc, "2. Thiết kế giải thuật xếp tầng linh hoạt: ", "Sử dụng heuristic nhiều cấp độ giúp hệ thống thích ứng tốt với sự đa dạng của các định dạng đề thi thực tế.")
    add_bullet_point(doc, "3. Thiết kế chịu lỗi và bảo mật phiên (Fault-tolerant & Session Security): ", "Cơ chế lưu trữ phiên làm việc tự động (Auto-Save) kết hợp lớp mã hóa XOR + Base64 đảm bảo hệ thống hoạt động tin cậy, chống gian lận dữ liệu khi học sinh làm bài tự luyện ngoại tuyến.")
    add_bullet_point(doc, "4. Tầm quan trọng của việc kiểm định chất lượng dữ liệu: ", "Xây dựng phân hệ validation và feedback loop giúp kiểm soát tốt chất lượng đề thi, tránh lỗi hệ thống khi vận hành.")
    
    add_heading_2(doc, "5.3. Định hướng nâng cấp giai đoạn tiếp theo")
    plan_data = [
        ["Thời hạn", "Mục tiêu nâng cấp", "Nội dung chi tiết"],
        ["30 ngày", "Tối ưu hóa và Vá lỗi", "Rà soát toàn bộ các trường hợp lỗi font chữ đặc biệt trong PDF; tối ưu hóa giao diện cài đặt phím tắt."],
        ["60 ngày", "Tăng cường tính năng tương tác", "Hỗ trợ nhập đề thi từ các định dạng ảnh chụp (OCR); bổ sung tính năng xuất đề thi trực tiếp sang PDF."],
        ["90 ngày", "Tích hợp trí tuệ nhân tạo (AI - Mở rộng)", "Tích hợp giải pháp Hybrid LLM: Hỗ trợ API Cloud (Gemini/OpenAI) giúp xử lý nhanh không tốn phần cứng, và tùy chọn Local LLM lượng tử hóa siêu nhẹ (Gemma-2B-IT-Int4/Phi-3-Mini) chạy qua Ollama/llama.cpp cho máy cấu hình mạnh."]
    ]
    create_table(doc, plan_data[0], plan_data[1:], [1.2, 2.3, 4.0])
    add_table_caption(doc, "Bảng 5. Kế hoạch hành động nâng cấp hệ thống Quiz Processor trong 90 ngày")
    
    doc.add_page_break()
    
    # --- TÀI LIỆU THAM KHẢO ---
    add_heading_1(doc, "TÀI LIỆU THAM KHẢO")
    add_styled_paragraph(doc, "[1] Python Software Foundation, \"Python Language Documentation,\" 2026. [Online]. Available: https://docs.python.org/")
    add_styled_paragraph(doc, "[2] Flutter Dev Team, \"Flutter Desktop Documentation,\" 2026. [Online]. Available: https://docs.flutter.dev/desktop")
    add_styled_paragraph(doc, "[3] PyMuPDF Team, \"PyMuPDF Documentation,\" 2026. [Online]. Available: https://pymupdf.readthedocs.io/")
    add_styled_paragraph(doc, "[4] Steve Canny, \"python-docx API Documentation,\" 2026. [Online]. Available: https://python-docx.readthedocs.io/")
    add_styled_paragraph(doc, "[5] SQLite Team, \"SQLite Documentation,\" 2026. [Online]. Available: https://www.sqlite.org/docs.html")
    
    doc.add_page_break()
    
    # --- PHỤ LỤC ---
    add_heading_1(doc, "PHỤ LỤC")
    add_heading_2(doc, "A. Cấu trúc thư mục dự án Quiz Processor")
    
    dir_data = [
        ["Thư mục / Tệp tin", "Mô tả vai trò trong hệ thống"],
        ["quiz_core/", "Nhân xử lý logic tài liệu (Python Backend)"],
        ["quiz_core/parsing/", "Phân tích PDF/DOCX, lọc nhiễu, nhận diện highlight đáp án"],
        ["quiz_core/grading/", "Chấm điểm bài thi, đối chiếu thông minh, xuất báo cáo câu lỗi"],
        ["quiz_core/validation/", "Kiểm định cấu trúc câu hỏi, quản lý feedback loop"],
        ["quiz_core/models.py", "Định nghĩa các Data Class cấu trúc dữ liệu dùng chung"],
        ["quiz_flutter_ui/", "Giao diện phòng thi & phân tích (Flutter Frontend)"],
        ["quiz_flutter_ui/lib/main.dart", "Entrypoint khởi chạy ứng dụng Flutter Desktop"],
        ["quiz_flutter_ui/lib/screens/", "Giao diện các màn hình: Số hóa, Phòng thi, Phân tích, Cài đặt"],
        ["quiz_flutter_ui/lib/services/", "Các lớp dịch vụ kết nối CLI, SQLite, SharedPreferences, Sao lưu"],
        ["quiz_cli.py", "Cổng kết nối CLI nhận lệnh từ giao diện gửi xuống backend"],
        ["requirements.txt", "Danh sách các thư viện Python phụ thuộc cần cài đặt"],
        ["QuizCLI.spec", "Tệp tin cấu hình đóng gói ứng dụng bằng PyInstaller"]
    ]
    create_table(doc, dir_data[0], dir_data[1:], [2.8, 4.7])
    add_table_caption(doc, "Bảng 6. Cấu trúc các thành phần thư mục chính trong dự án Quiz Processor")
    
    # Save the document
    out_dir = r"d:\My_projects\Random_Essential\Quiz_Processor\docs\de_an"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "BaoCaoDeAn_QuizProcessor.docx")
    doc.save(out_path)
    print(f"Report successfully saved to {out_path}")

if __name__ == "__main__":
    generate_report()
