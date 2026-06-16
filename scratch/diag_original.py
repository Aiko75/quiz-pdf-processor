"""
Diagnostic tren file GOC (ORIGINAL) de xac nhan highlight co that su ton tai
"""
import docx
from pathlib import Path
from lxml import etree
import re

doc_path = Path('docs/test/trac-nghiem-1-chu-nghia-xa-hoi-khoa-hoc-cnxhkh_co_dap_an_ORIGINAL.docx')
doc = docx.Document(doc_path)

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
Q_PATTERN = re.compile(r'^Câu\s+\d+[\.:)]?\s')

def get_run_highlight_info(run):
    rpr = run._r.find(f'{{{NS}}}rPr')
    if rpr is None:
        return None, None
    hl = rpr.find(f'{{{NS}}}highlight')
    shd = rpr.find(f'{{{NS}}}shd')
    hl_val = hl.get(f'{{{NS}}}val') if hl is not None else None
    shd_fill = None
    if shd is not None:
        fill = shd.get(f'{{{NS}}}fill')
        val = shd.get(f'{{{NS}}}val', '')
        if fill and fill not in ('auto', '000000', 'FFFFFF', 'auto') and val not in ('clear', 'none', ''):
            shd_fill = fill
    return hl_val, shd_fill

print(f'[INFO] File: {doc_path}')
print(f'[INFO] Total paragraphs: {len(doc.paragraphs)}')
print()

# Kiem tra 3 cau dau
print('=== 3 CAU HOI DAU TIEN ===')
found = 0
for i, para in enumerate(doc.paragraphs[:50]):
    t = para.text.strip()
    if Q_PATTERN.match(t):
        found += 1
        if found > 3:
            break
        print(f'\n[Q{found}] para[{i}]: {t[:80]}')
        for j in range(i+1, min(i+8, len(doc.paragraphs))):
            pt = doc.paragraphs[j].text.strip()
            if Q_PATTERN.match(pt):
                break
            if not pt:
                continue
            highlights = []
            for run in doc.paragraphs[j].runs:
                hl, shd = get_run_highlight_info(run)
                if hl:
                    highlights.append(f'highlight={hl}')
                if shd:
                    highlights.append(f'shd={shd}')
            marker = str(highlights) if highlights else '(no highlight)'
            print(f'   [{j}] {marker:40s} {pt[:60]}')

# Full count
print('\n=== TONG KET TOAN BO FILE GOC ===')
q_hl = []
q_shd = []
q_none = []

i = 0
while i < len(doc.paragraphs):
    t = doc.paragraphs[i].text.strip()
    if Q_PATTERN.match(t):
        found_hl, found_shd = False, False
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            pt = doc.paragraphs[j].text.strip()
            if Q_PATTERN.match(pt):
                break
            for run in doc.paragraphs[j].runs:
                hl, shd = get_run_highlight_info(run)
                if hl:
                    found_hl = True
                if shd:
                    found_shd = True
        if found_hl:
            q_hl.append((i, t[:60]))
        elif found_shd:
            q_shd.append((i, t[:60]))
        else:
            q_none.append((i, t[:60]))
    i += 1

print(f'Cau co w:highlight: {len(q_hl)}')
print(f'Cau co w:shd:       {len(q_shd)}')
print(f'Cau khong hl gi:    {len(q_none)}')
print()

# Dump XML cua 1 option de xac nhan
print('=== XML RAW OPTION PARA (para[1] cua file goc) ===')
p = doc.paragraphs[1]
print(f'Text: "{p.text[:80]}"')
print(etree.tostring(p._p, pretty_print=True).decode('utf-8')[:1200])
