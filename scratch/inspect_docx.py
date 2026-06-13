import docx
from pathlib import Path

doc_path = Path("docs/test/trac-nghiem-1-chu-nghia-xa-hoi-khoa-hoc-cnxhkh_co_dap_an.docx")
doc = docx.Document(doc_path)

q_nums = [38, 40, 59, 68, 77, 81]

print("=== INSPECTING DOCX PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text: continue
    
    # Check if this paragraph contains a question prefix for our target questions
    for q in q_nums:
        if text.startswith(f"Câu {q}:") or text.startswith(f"Câu {q}.") or text.startswith(f"{q}."):
            print(f"\n--- Paragraphs around Question {q} (Index {i}) ---")
            # Print 5 paragraphs before and 5 after
            start = max(0, i - 1)
            end = min(len(doc.paragraphs), i + 10)
            for j in range(start, end):
                p_text = doc.paragraphs[j].text.strip()
                highlighted = any(run.font.highlight_color is not None for run in doc.paragraphs[j].runs)
                bold = any(run.bold for run in doc.paragraphs[j].runs)
                print(f"[{j}] {'(H)' if highlighted else ''}{'(B)' if bold else ''} {p_text}")
