"""
Diagnostic: Kiem tra cac loai highlight trong DOCX
- run.font.highlight_color -> w:highlight element
- w:shd element (background shading) -> python-docx KHONG tu dong doc
"""
import docx
from pathlib import Path
from lxml import etree
import re

doc_path = Path('docs/test/trac-nghiem-1-chu-nghia-xa-hoi-khoa-hoc-cnxhkh_co_dap_an.docx')
doc = docx.Document(doc_path)

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def check_run_highlight(run):
    """Kiem tra ca 2 kieu highlight: w:highlight va w:shd"""
    rpr = run._r.find(f'{{{NS}}}rPr')
    if rpr is None:
        return None, None
    
    # Kieu 1: w:highlight (standard highlight - python-docx doc duoc)
    hl = rpr.find(f'{{{NS}}}highlight')
    hl_val = hl.get(f'{{{NS}}}val') if hl is not None else None
    
    # Kieu 2: w:shd (background shading - python-docx KHONG doc duoc)
    shd = rpr.find(f'{{{NS}}}shd')
    shd_fill = None
    if shd is not None:
        shd_fill = shd.get(f'{{{NS}}}fill')
        shd_val = shd.get(f'{{{NS}}}val', '')
        # Bo qua shd vo hieu (auto/none/000000/FFFFFF)
        if shd_fill in ('auto', '000000', 'FFFFFF', 'auto', None) or shd_val in ('clear', 'none'):
            shd_fill = None
    
    return hl_val, shd_fill

def para_has_any_highlight(para):
    """Kiem tra ca 2 loai highlight trong tat ca runs cua paragraph."""
    for run in para.runs:
        hl, shd = check_run_highlight(run)
        if hl or shd:
            return True, hl, shd
    return False, None, None

Q_PATTERN = re.compile(r'^Câu\s+\d+[\.:)]?\s')

print('=== KIEM TRA CHI TIET HIGHLIGHT (ca w:highlight va w:shd) ===\n')

# Kiem tra 20 cau dau
print('--- 10 PARAGRAPH CAU HOI DAU TIEN (chi thi options) ---')
found_q = 0
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if Q_PATTERN.match(t):
        found_q += 1
        if found_q > 3:
            break
        print(f'\n[{i}] {t[:80]}')
        # Kiem tra 6 paragraph sau
        for j in range(i+1, min(i+8, len(doc.paragraphs))):
            pt = doc.paragraphs[j].text.strip()
            if Q_PATTERN.match(pt):
                break
            if not pt:
                continue
            has_hl, hl_val, shd_val = para_has_any_highlight(doc.paragraphs[j])
            marker = ''
            if hl_val:
                marker = f'[w:highlight={hl_val}]'
            elif shd_val:
                marker = f'[w:shd fill={shd_val}]'
            print(f'   [{j}] {marker:30s} {pt[:70]}')

print('\n\n--- TONG KET TOAN BO FILE ---')
q_with_hl = []
q_with_shd = []
q_none = []
i = 0
while i < len(doc.paragraphs):
    t = doc.paragraphs[i].text.strip()
    if Q_PATTERN.match(t):
        q_text = t[:60]
        found_hl = False
        found_shd = False
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            pt = doc.paragraphs[j].text.strip()
            if Q_PATTERN.match(pt):
                break
            has_hl, hl_val, shd_val = para_has_any_highlight(doc.paragraphs[j])
            if hl_val:
                found_hl = True
            if shd_val:
                found_shd = True
        if found_hl:
            q_with_hl.append((i, q_text))
        elif found_shd:
            q_with_shd.append((i, q_text))
        else:
            q_none.append((i, q_text))
    i += 1

print(f'Cau co w:highlight (standard):  {len(q_with_hl)}')
print(f'Cau co w:shd (background shad): {len(q_with_shd)}')
print(f'Cau KHONG co bat ky highlight:   {len(q_none)}')

if q_with_shd:
    print('\nCac cau dung w:shd:')
    for idx, text in q_with_shd[:10]:
        print(f'  [{idx}] {text}')

if q_none[:5]:
    print('\nMau cac cau khong co highlight:')
    for idx, text in q_none[:5]:
        print(f'  [{idx}] {text}')

# Dump raw XML cua option dau tien de xem cau truc
print('\n\n--- RAW XML OPTION DAU TIEN ---')
for i, para in enumerate(doc.paragraphs[:10]):
    t = para.text.strip()
    if t and (t.startswith('A.') or t.startswith('B.')):
        print(f'Paragraph [{i}]: "{t[:60]}"')
        print(etree.tostring(para._p, pretty_print=True).decode('utf-8')[:800])
        print()
        break
