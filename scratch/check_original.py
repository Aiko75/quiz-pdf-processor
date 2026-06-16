"""
Kiem tra nhanh file ORIGINAL: watermark, so cau, cau 36, cau 38/40/59/68/77/81
"""
import docx
from pathlib import Path
from lxml import etree
import re

doc_path = Path('docs/test/trac-nghiem-1-chu-nghia-xa-hoi-khoa-hoc-cnxhkh_co_dap_an_ORIGINAL.docx')
doc = docx.Document(doc_path)

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
Q_PATTERN = re.compile(r'^Câu\s+\d+[\.:)]?\s')

WM_PATTERNS = ['messages.pdf_cover_qr_code_label', 'messages.downloaded_by', 'lOMoARcPSD']

def has_highlight(para):
    for run in para.runs:
        rpr = run._r.find(f'{{{NS}}}rPr')
        if rpr is None: continue
        hl = rpr.find(f'{{{NS}}}highlight')
        if hl is not None: return True
    return False

print(f'Total paragraphs: {len(doc.paragraphs)}')

# 1. Kiem tra watermark trong file ORIGINAL
print('\n=== WATERMARK TRONG FILE ORIGINAL ===')
wm_found = []
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    for wm in WM_PATTERNS:
        if wm.lower() in t.lower():
            wm_found.append((i, t[:120]))
            break
if wm_found:
    print(f'[!] Co {len(wm_found)} doan co watermark:')
    for idx, text in wm_found:
        print(f'  [{idx}] {text}')
else:
    print('[OK] Khong co watermark nao trong file ORIGINAL')

# 2. Kiem tra cac cau trong feedback (36, 38, 40, 59, 68, 77, 81)
print('\n=== KIEM TRA CAC CAU TRONG FEEDBACK ===')
check_q = ['Câu 36:', 'Câu 38:', 'Câu 40:', 'Câu 59:', 'Câu 68:', 'Câu 77:', 'Câu 81:']
for anchor in check_q:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(anchor):
            print(f'\n[{i}] {para.text.strip()[:90]}')
            for j in range(i+1, min(i+8, len(doc.paragraphs))):
                pt = doc.paragraphs[j].text.strip()
                if Q_PATTERN.match(pt): break
                if not pt: continue
                hl = has_highlight(doc.paragraphs[j])
                print(f'   [{j}] {"[H]" if hl else "   "} {pt[:80]}')
            break
    else:
        print(f'\n[NOT FOUND] {anchor}')

# 3. Cac cau khong co highlight
print('\n=== TONG KET ===')
no_hl = []
i = 0
while i < len(doc.paragraphs):
    t = doc.paragraphs[i].text.strip()
    if Q_PATTERN.match(t):
        found = any(has_highlight(doc.paragraphs[j])
                    for j in range(i+1, min(i+10, len(doc.paragraphs)))
                    if not Q_PATTERN.match(doc.paragraphs[j].text.strip()))
        if not found:
            no_hl.append((i, t[:70]))
    i += 1
print(f'Cau khong co highlight: {len(no_hl)}')
for idx, text in no_hl:
    print(f'  [{idx}] {text}')
