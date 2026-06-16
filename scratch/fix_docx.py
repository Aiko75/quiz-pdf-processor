"""
fix_docx_v3.py
==============
Phiên bản 3: Bổ sung thêm:
  - Xóa watermark bị nhúng trong nội dung câu hỏi (messages.downloaded_by, etc.)
  - Lưu ra file MỚI thay vì ghi đè (suffix _fixed)
  - In đường dẫn file kết quả rõ ràng

Watermark pattern phổ biến cứ mỗi 7 câu (câu 1, 8, 16, 23, 30, 36, 43, 50...):
  "messages.pdf_cover_qr_code_label <tên file> <tên file> messages.downloaded_by"
  "messages.downloaded_by <số câu>. <nội dung câu>"
"""

import docx
import re
from pathlib import Path
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import copy

# File GOC co highlight day du (copy tu D:\Download\cnxhkh\)
DOCX_PATH = Path("docs/test/trac-nghiem-1-chu-nghia-xa-hoi-khoa-hoc-cnxhkh_co_dap_an_ORIGINAL.docx")

# =====================================================================
# Du lieu sua doi theo cau hoi (xac nhan tu file ORIGINAL)
# Cau 36: DA CO highlight (B. Bon) -> KHONG can sua
# Cau 38: A/B/C/D trong -> dien noi dung, highlight B
# Cau 40: A co hl, B/C trong -> dien noi dung
# Cau 59: A/B trong, C sai ("C. Mac & Ph.Ang ghen" thuc ra la B), thieu D
# Cau 68: A/B trong -> dien noi dung, highlight B
# Cau 77: A/B trong, C la B sai vi tri, thieu D -> sua toan bo
# Cau 81: KHONG TON TAI trong file goc
# =====================================================================
QUESTION_FIXES = [
    {
        "question_anchor": "Câu 38:",
        "options": [
            ("A. ", "A. 3", False),
            ("B. ", "B. 4", True),
            ("C. ", "C. 5", False),
            ("D. ", "D. 6", False),
        ],
        "insert_after": None
    },
    {
        "question_anchor": "Câu 40:",
        "options": [
            # A. da co highlight -> giu nguyen, chi dam bao text dung
            ("A. Cả ba đáp án trên", "A. Cả ba đáp án trên", True),
            ("B. ", "B. Những sai lầm của Đảng và của những người lãnh đạo cấp cao nhất Đảng Cộng sản Liên Xô.", False),
            ("C. ", "C. Quan niệm và vận dụng không đúng đắn về CNXH", False),
        ],
        "insert_after": None
    },
    {
        "question_anchor": "Câu 59:",
        "options": [
            ("A. ", "A. C.Mác", False),
            ("B. ", "B. C.Mác & Ph.Ăng ghen", False),
            # C hien tai = "C. Mac & Ph.Ang ghen" -> sua thanh "C. Ho Chi Minh"
            ("C. Mác", "C. Hồ Chí Minh", False),
        ],
        "insert_after": {
            "after_anchor": "C. Hồ Chí Minh",
            "new_text": "D. V.I Lênin",
            "is_correct": True
        }
    },
    {
        "question_anchor": "Câu 68:",
        "options": [
            ("A. ", "A. Sai", False),
            ("B. ", "B. Đúng", True),
        ],
        "insert_after": None
    },
    {
        "question_anchor": "Câu 77:",
        "options": [
            ("A. ", "A. C.Mác", False),
            ("B. ", "B. C.Mác & Ph.Ăng ghen", True),
            # C hien tai = "C. Mac & Ph.Ang ghen" (sai) -> sua thanh C. Ph.Ang ghen
            ("C. Mác", "C. Ph.Ăng ghen", False),
        ],
        "insert_after": {
            "after_anchor": "C. Ph.Ăng ghen",
            "new_text": "D. V.I Lênin.",
            "is_correct": False
        }
    }
]

# Khong can xoa watermark vi file ORIGINAL da sach
WATERMARK_STANDALONE = None
# Output file (khong ghi de len ORIGINAL)
OUTPUT_PATH = Path("docs/test/trac-nghiem-1-chu-nghia-xa-hoi-khoa-hoc-cnxhkh_co_dap_an_FIXED.docx")

# Watermark patterns nhúng trong nội dung câu hỏi
INLINE_WATERMARK_PATTERNS = [
    r'messages\.pdf_cover_qr_code_label[^\n]*',   # tiêu đề file lặp lại
    r'messages\.downloaded_by\s*',                 # prefix watermark
    r'lOMoARcPSD\|[^\s]*\s*',                      # mã hóa Studocu
]

# Compile pattern tổng hợp
COMBINED_WM_RE = re.compile('|'.join(INLINE_WATERMARK_PATTERNS), re.IGNORECASE)

# Pattern nhan dien so cau trong watermark (vi watermark chua so cau: "8. Ai da viet...")
# Sau khi xoa watermark, cau hoi bi du so thu tu: "Câu 8: messages.downloaded_by 8. Ai..."
DUPLICATE_Q_NUMBER_RE = re.compile(r'^(\d+)\.\s+')  # prefix "8. " sau watermark


def clean_watermark_from_text(text: str) -> str:
    """Xóa watermark inline khỏi text câu hỏi, giữ lại nội dung thực."""
    cleaned = COMBINED_WM_RE.sub('', text)
    # Xóa số câu trùng lặp nếu có (vd: "8. Ai da viet..." -> "Ai da viet...")
    cleaned = DUPLICATE_Q_NUMBER_RE.sub('', cleaned.strip())
    return cleaned.strip()


def find_paragraph_index(doc, text_starts_with: str, start_from: int = 0) -> int:
    for i in range(start_from, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith(text_starts_with):
            return i
    return -1


def set_paragraph_text(para, text: str, is_correct: bool):
    p = para._p
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r in p.findall(f'{{{ns}}}r'):
        p.remove(r)
    for hl in p.findall(f'{{{ns}}}hyperlink'):
        p.remove(hl)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if is_correct:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    print(f"    [SET] '{text}' | highlight={is_correct}")


def insert_paragraph_after(doc, ref_para, text: str, is_correct: bool):
    ref_p = ref_para._p
    new_p = copy.deepcopy(ref_p)
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r in new_p.findall(f'{{{ns}}}r'):
        new_p.remove(r)
    for hl in new_p.findall(f'{{{ns}}}hyperlink'):
        new_p.remove(hl)
    ref_p.addnext(new_p)
    new_para = None
    for para in doc.paragraphs:
        if para._p is new_p:
            new_para = para
            break
    if new_para:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if is_correct:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        print(f"    [INSERT] '{text}' | highlight={is_correct}")
    else:
        print(f"    [ERROR] Không thể chèn paragraph mới")


def delete_paragraph(para):
    text_preview = para.text[:80]
    p = para._p
    p.getparent().remove(p)
    print(f"    [DELETE] '{text_preview}'")


def main():
    print(f"[LOAD] {DOCX_PATH}")
    doc = docx.Document(DOCX_PATH)
    print(f"[INFO] Tổng số paragraph: {len(doc.paragraphs)}")

    # ===== BƯỚC 0: Xóa watermark nhúng trong nội dung câu hỏi =====
    print("\n=== BƯỚC 0: DỌN SẠCH WATERMARK INLINE ===")
    q_pattern = re.compile(r'^Câu\s+\d+[\.:)]?\s')
    wm_cleaned = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if q_pattern.match(t) and COMBINED_WM_RE.search(t):
            # Tách phần "Câu X:" ra
            q_prefix_match = re.match(r'^(Câu\s+\d+:?\s*)', t)
            q_prefix = q_prefix_match.group(1) if q_prefix_match else ''
            rest = t[len(q_prefix):]
            cleaned_rest = clean_watermark_from_text(rest)
            new_text = q_prefix + cleaned_rest

            # Ghi lại run
            p = para._p
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for r in p.findall(f'{{{ns}}}r'):
                p.remove(r)
            for hl in p.findall(f'{{{ns}}}hyperlink'):
                p.remove(hl)
            run = para.add_run(new_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

            print(f"  [CLEAN] [{para._p.getparent().index(para._p)}] '{t[:60]}...'")
            print(f"      --> '{new_text[:80]}'")
            wm_cleaned += 1

    print(f"[OK] Đã dọn sạch {wm_cleaned} câu hỏi có watermark inline.\n")

    # ===== BƯỚC 1: Sửa đổi từng câu hỏi trong QUESTION_FIXES =====
    print("=== BƯỚC 1: SỬA ĐỔI CÂU HỎI TRONG FEEDBACK ===")
    for fix in QUESTION_FIXES:
        anchor = fix["question_anchor"]
        q_idx = find_paragraph_index(doc, anchor)

        if q_idx == -1:
            print(f"\n[SKIP] Không tìm thấy: '{anchor}'")
            continue

        print(f"\n[QUESTION] Câu tìm thấy tại [{q_idx}]: {doc.paragraphs[q_idx].text.strip()[:60]}")

        for option_prefix, new_text, is_correct in fix["options"]:
            found = False
            # Normalize: strip trailing space to match both "A." and "A. text"
            prefix_stripped = option_prefix.rstrip()
            for search_idx in range(q_idx + 1, min(q_idx + 12, len(doc.paragraphs))):
                p_text = doc.paragraphs[search_idx].text.strip()
                if p_text and (p_text.startswith("Câu ") and ":" in p_text[:10]):
                    break
                # Match "A." (empty option) OR "A. text..." OR full prefix
                if (p_text == prefix_stripped or
                    p_text.startswith(option_prefix) or
                    p_text.startswith(prefix_stripped + '.')):
                    set_paragraph_text(doc.paragraphs[search_idx], new_text, is_correct)
                    found = True
                    break
            if not found:
                print(f"    [NOT FOUND] Đáp án với prefix: '{option_prefix}'")


        if fix["insert_after"]:
            ins = fix["insert_after"]
            ref_idx = find_paragraph_index(doc, ins["after_anchor"], q_idx + 1)
            if ref_idx == -1:
                print(f"    [INSERT-SKIP] Không tìm thấy anchor: '{ins['after_anchor']}'")
            else:
                next_texts = [doc.paragraphs[i].text.strip() for i in range(ref_idx + 1, min(ref_idx + 3, len(doc.paragraphs)))]
                if any(t.startswith("D. ") for t in next_texts):
                    for i in range(ref_idx + 1, min(ref_idx + 3, len(doc.paragraphs))):
                        if doc.paragraphs[i].text.strip().startswith("D. "):
                            set_paragraph_text(doc.paragraphs[i], ins["new_text"], ins["is_correct"])
                            break
                else:
                    insert_paragraph_after(doc, doc.paragraphs[ref_idx], ins["new_text"], ins["is_correct"])

    # ===== BƯỚC 2: Xóa watermark paragraph độc lập (Câu 81) =====
    print(f"\n=== BƯỚC 2: XÓA WATERMARK PARAGRAPH ĐỘC LẬP ===")
    if WATERMARK_STANDALONE is None:
        print(f"    [SKIP] WATERMARK_STANDALONE = None -> file goc da sach, bo qua buoc nay.")
    else:
        wm_idx = find_paragraph_index(doc, WATERMARK_STANDALONE)
        if wm_idx != -1:
            delete_paragraph(doc.paragraphs[wm_idx])
        else:
            print(f"    [SKIP] Không tìm thấy '{WATERMARK_STANDALONE}' (đã xóa trước đó)")


    # ===== BƯỚC 3: Kiểm tra sau khi sửa =====
    print(f"\n=== BƯỚC 3: KIỂM TRA SAU KHI SỬA ===")
    check_anchors = ["Câu 1:", "Câu 8:", "Câu 16:", "Câu 38:", "Câu 40:", "Câu 59:", "Câu 68:", "Câu 77:", "Câu 81:"]
    for anchor in check_anchors:
        idx = find_paragraph_index(doc, anchor)
        if idx == -1:
            print(f"\n  [{anchor}] Không tìm thấy")
            continue
        print(f"\n  [{idx}] {doc.paragraphs[idx].text.strip()[:90]}")
        for j in range(idx + 1, min(idx + 7, len(doc.paragraphs))):
            pt = doc.paragraphs[j].text.strip()
            if not pt:
                continue
            if pt.startswith("Câu ") and ":" in pt[:12]:
                break
            highlighted = any(r.font.highlight_color is not None for r in doc.paragraphs[j].runs)
            marker = "✓(H)" if highlighted else "    "
            print(f"       [{j}] {marker} {pt[:80]}")

    # ===== BƯỚC 4: Lưu file MỚI (không ghi đè ORIGINAL) =====
    out_path = OUTPUT_PATH
    doc.save(out_path)
    print(f"\n{'='*60}")
    print(f"[DONE] File đã sửa được lưu tại:")
    print(f"  >>> {out_path.resolve()} <<<")
    print(f"{'='*60}")
    print(f"  Để sử dụng: thay thế đường dẫn file đáp án thành file trên.")
    print(f"  File gốc KHÔNG bị thay đổi: {DOCX_PATH.resolve()}")


if __name__ == "__main__":
    main()
