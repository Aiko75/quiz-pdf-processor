import re
from pathlib import Path
from typing import List, Dict, Any
import docx
from docx.shared import Pt
from .patterns import QUESTION_PATTERN, QUESTION_PREFIX_ONLY_PATTERN
from .core import match_option_line
from .utils import normalize_question_text, extract_logical_index
from ..models import QuestionData, QuizQuestionState, QuizOptionState

def write_output_files(questions: List[QuestionData], input_file: Path, output_dir: Path) -> None:
    base_name = input_file.stem
    with_answers = output_dir / f"{base_name}_co_dap_an.docx"
    practice = output_dir / f"{base_name}_de_lam.docx"
    ans_doc = docx.Document()
    practice_doc = docx.Document()
    for doc in (ans_doc, practice_doc):
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
    for idx, q in enumerate(questions, start=1):
        norm_q = re.sub(r"^(?:Câu\s*\d+[\.:\)]?|Question\s*\d+[\.:\)]?|\d+[\.)])\s*", "", q.question, flags=re.IGNORECASE).strip()
        ans_doc.add_paragraph(f"Câu {idx}: {norm_q}")
        practice_doc.add_paragraph(f"Câu {idx}: {norm_q}")
        for opt in q.options:
            line = f"{opt.label}. {opt.text}"
            ans_p = ans_doc.add_paragraph()
            ans_r = ans_p.add_run(line)
            if q.answer_label == opt.label:
                ans_r.bold = True
                ans_r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            practice_doc.add_paragraph(line)
        ans_doc.add_paragraph("")
        practice_doc.add_paragraph("")
    ans_doc.save(with_answers)
    practice_doc.save(practice)

def parse_docx_questions(docx_path: Path) -> List[Dict[str, Any]]:
    doc = docx.Document(docx_path)
    questions = []
    current_q = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        if QUESTION_PATTERN.match(text) or QUESTION_PREFIX_ONLY_PATTERN.match(text):
            if current_q: questions.append(current_q)
            current_q = {"question": text, "options": {}, "correct_answer": None}
            continue
            
        opt_match = match_option_line(text)
        if opt_match and current_q:
            label = opt_match.group(1).upper()
            content = opt_match.group(2).strip()
            current_q['options'][label] = content
            
            # Simple heuristic for correct answer: bold
            if any(run.bold for run in para.runs):
                current_q['correct_answer'] = label
        elif current_q and not current_q['options']:
            current_q['question'] += " " + text
            
    if current_q: questions.append(current_q)
    
    # Filter questions
    valid_questions = []
    for q in questions:
        q['question'] = normalize_question_text(q['question'])
        if len(q['question']) < 5 or q['question'].isdigit():
            continue
        # Check if it has at least 2 options and all options are reasonably short (< 500 chars)
        if len(q['options']) >= 2 and all(len(opt_text) < 500 for opt_text in q['options'].values()):
            valid_questions.append(q)
            
    total_potential = len([q for q in questions if len(normalize_question_text(q['question'])) >= 5 and not normalize_question_text(q['question']).isdigit()])
    ratio = len(valid_questions) / total_potential if total_potential > 0 else 0.0
    if ratio < 0.6:
        return []
        
    return valid_questions

def parse_docx_questions_for_grading(docx_path: Path) -> List[QuizQuestionState]:
    doc = docx.Document(docx_path)
    questions = []
    current_q = None
    
    def finalize_current_q(q):
        if not q: return
        
        # Priority 1: Yellow Highlight (most reliable)
        highlighted = [l for l, opt in q.options.items() if opt.highlighted]
        if highlighted:
            q.highlighted_labels = highlighted
            return
            
        # Priority 2: Bold (only if exactly one is bold)
        bolded = [l for l, opt in q.options.items() if opt.is_bold]
        if len(bolded) == 1:
            q.highlighted_labels = bolded
            q.options[bolded[0]].highlighted = True
            return
            
        # Priority 3: Colored (non-black)
        colored = [l for l, opt in q.options.items() if opt.color_rgb and opt.color_rgb != (0, 0, 0)]
        if len(colored) == 1:
            q.highlighted_labels = colored
            q.options[colored[0]].highlighted = True
            return

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        if QUESTION_PATTERN.match(text) or QUESTION_PREFIX_ONLY_PATTERN.match(text):
            if current_q:
                finalize_current_q(current_q)
                questions.append(current_q)
            l_idx = extract_logical_index(text)
            current_q = QuizQuestionState(question=text, options={}, highlighted_labels=[], logical_index=l_idx)
            continue
            
        opt_match = match_option_line(text)
        if opt_match and current_q:
            label = opt_match.group(1).upper()
            content = opt_match.group(2).strip()
            
            # Check for actual highlight color
            is_highlighted = any(run.font.highlight_color is not None for run in para.runs)
            is_bold = any(run.bold for run in para.runs)
            
            color_rgb = None
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    color_rgb = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                    if color_rgb != (0, 0, 0):
                        # Treat non-black text as a marker too
                        pass 
                    break
            
            current_q.options[label] = QuizOptionState(
                label=label,
                text=content,
                highlighted=is_highlighted,
                is_bold=is_bold,
                color_rgb=color_rgb
            )
        elif current_q and not current_q.options:
            current_q.question += " " + text
            
    if current_q:
        finalize_current_q(current_q)
        questions.append(current_q)
        
    # Filter questions
    valid_questions = []
    for q in questions:
        q.question = normalize_question_text(q.question)
        if len(q.question) < 5 or q.question.isdigit():
            continue
        # Check if it has at least 2 options and all options are reasonably short (< 500 chars)
        if len(q.options) >= 2 and all(len(opt.text) < 500 for opt in q.options.values()):
            valid_questions.append(q)
            
    total_potential = len([q for q in questions if len(normalize_question_text(q.question)) >= 5 and not normalize_question_text(q.question).isdigit()])
    ratio = len(valid_questions) / total_potential if total_potential > 0 else 0.0
    if ratio < 0.6:
        return []
        
    return valid_questions
